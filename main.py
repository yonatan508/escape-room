from tkinter import *
from tkinter import filedialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import docx

root = Tk()
root.title("פרויקט")
root.geometry("1000x1000")

canvas = Canvas(root, width=1000, height=1000)
canvas.pack()
def debugger():
    if MemoryError: raise Exception('There is an error with the memory')
    elif EOFError: raise Exception('There is an error with the code')
    elif EnvironmentError: raise Exception('There is an error with the modules or the lang may be corrupted')
    else: print('you fucked up the program')

def timer(tk, time_label, no):
    time_label.configure(text=now)
    tk.after(1000, timer)
    now += 1

def story():
    file_path = askdirectory()
    docx1 = docx.Document()

    docx1.add_heading("escape room story",0)
    docx_par = docx1.add_paragraph("The story of the game")
    docx_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    docx.add_run("""
        שלום, אתם בשנת 2174 וכל העולם נדבק בוירוס בשם LCA-100, רק אתם יכולים להציל אותו, אבל, תצטרכו לפתור את החידות הבאות בשביל להציל את העולם, חידות אלו הן חידות שאפילו אני לא הצלחתי לפתור והם יעזרו לכם להגיע למקור המחלה ולהציל את האנושות, אני בניתי מכונת זמן אשר תיקח אתכם ליעד הראשון, 1 למאי, 2028, שם נמצא האדם הראשון לחלות במחלה, אתם תצטרכו לאתר אותו בשביל למנוע ממנו לראות אנשים אחרים במשך חודש, אני לא יכול לספר לכם יותר מזה.
את השאר אתם תצטרכו להבין לבד, מבחינת הזמן שלכם, יש לכם חצי שעה לפתור את החידות.
בהצלחה בדרך, רמז: לכל אחד מהאנשים שתמצאו יש מחלה שונה בקצת מהשני. 
""")
    docx.add_run("בהצלחה")
    os.chdir(file_path)
    docx.save(file_path)


def word_download(clue_bool,answer_bool, level, number_levels=None):
    if clue_bool and answer_bool: raise Exception('You have made a mistake with the menus command')
    saving_path = filedialog.askdirectory()
    print(saving_path)
    if len(number_levels) == 1: heading = f"level {level}"
    else: heading = "all the levels" 
    docx1 = docx.Document()

    if clue_bool: docx1.add_heading("the clues of {}".format(heading),0)
    elif not clue_bool and heading[-1] == "s": 
            docx1.add_heading("questions and answers of {}".format(heading),0)
    else: docx1.add_heading("question and answer of {}".format(heading),0)

    list = ""
    
    if len(number_levels) > 1:
        for number in number_levels:
            list += f"{number}, "
        list = list[0:-1:1]
    if clue_bool and len(number_levels) == 1: doc_par = docx1.add_paragraph("The clues of question number {}:".format(level))
    elif clue_bool and len(number_levels) > 1: doc_par = docx1.add_paragraph("The clues of questions number {}:".format(list))
    elif not clue_bool and len(number_levels) == 1: doc_par = docx1.add_paragraph("the question {} answer and story:".format(level))
    else: doc_par = docx1.add_paragraph("the questions {} answers and stories:".format(list))
    doc_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if len(number_levels) == 1 and clue_bool and not answer_bool and level == 1: 
        doc_par.add_run("""
הרמזים של השאלה הראשונה:
שתמש בטבלת המחזורית  .1
חשוב איזה סוג חומר יכול להישרף  .2
חשוב מה הופך חומר מסוים לאורגני  .3
בר את המספרים האטומים שמאפיינים חומר אורגני  .4
""")

    if len(number_levels) == 1 and clue_bool and not answer_bool and level == 2: 
        doc_par.add_run("""
הרמזים של השאלה השנייה:
תשתמש בטבלה המחזורית  .1
תכתוב את המספרים הטומיים של כל יסוד  .2
תשתמש בגימטריה  .3
""")

    if len(number_levels) == 1 and clue_bool and not answer_bool and level == 3: 
        doc_par.add_run("""
הרמזים של השאלה השלישית: 
חומר אורגני = 1 וחומר לא אורגני = 0  .1
תמירו את המספר שהיתקבל למספר רגיל ממספר בינארי בכל אחד מהמשפטים  .2
תכתבו את המספרים אחד אחרי השני בסדר של המספרים כך המספר הראשון יהיה התוצאה של המשפט הגבוהה יותר  .3
טבלת המרה:
המספר 2 – 10
המספר 4 – 100
המספר 6 – 110
המספר 8 – 1000
המספר 10 – 1010
המספר 12 – 1100
המספר 14 - 1110

""")

    if len(number_levels) == 1 and clue_bool and not answer_bool and level == 4: 
        doc_par.add_run("""
הרמזים של השאלה הרביעית:
תזכרו איזה אברונים נמצאים בכל תא ומה כל אחד עושה  .1
תחפסו אברון שאין אותו לכל היצורים החיים  .2
""")

    if len(number_levels) == 1 and clue_bool and not answer_bool and level == 5: 
        doc_par.add_run("""
הרמזים של השאלה החמישית:
תשתמשו בטבלה המחזורית  .1
תמירו כל אות או סמל כימי למספר האטומי שלו  .2
""")

    if len(number_levels) > 1 and clue_bool and not answer_bool: 
        doc_par.add_run("""
הרמזים של השאלה הראשונה:
שתמש בטבלת המחזורית  .1
חשוב איזה סוג חומר יכול להישרף  .2
חשוב מה הופך חומר מסוים לאורגני  .3
בר את המספרים האטומים שמאפיינים חומר אורגני  .4

הרמזים של השאלה השנייה:
תשתמש בטבלה המחזורית  .1
תכתוב את המספרים הטומיים של כל יסוד  .2
תשתמש בגימטריה  .3

הרמזים של השאלה השלישית: 
חומר אורגני = 1 וחומר לא אורגני = 0  .1
תמירו את המספר שהיתקבל למספר רגיל ממספר בינארי בכל אחד מהמשפטים  .2
תכתבו את המספרים אחד אחרי השני בסדר של המספרים כך המספר הראשון יהיה התוצאה של המשפט הגבוהה יותר  .3
טבלת המרה:
המספר 2 – 10
המספר 4 – 100
המספר 6 – 110
המספר 8 – 1000
המספר 10 – 1010
המספר 12 – 1100
המספר 14 - 1110


הרמזים של השאלה הרביעית:
תזכרו איזה אברונים נמצאים בכל תא ומה כל אחד עושה  .1
תחפסו אברון שאין אותו לכל היצורים החיים  .2

הרמזים של השאלה החמישית:
תשתמשו בטבלה המחזורית  .1
תמירו כל אות או סמל כימי למספר האטומי שלו  .2
""")
    if len(number_levels) == 1 and not clue_bool and level == 1:
        docx1.add_run("""

            """)
    if saving_path in [" ", "", None]: pass
    else: docx1.save(f"{saving_path}/escape room.docx")

def questions(tuple):#(title, level, question, sizepx, size, xx, yy, x, y, ex, ey, result)):
    title = tuple[0]
    level = tuple[1]
    question = tuple[2]
    sizepx = tuple[3]
    size = tuple[4]
    xx = tuple[5]
    yy = tuple[6]
    x = tuple[7]
    y = tuple[8]
    ex = tuple[9]
    ey = tuple[10]
    result = tuple[11]

    master = Tk()
    master.configure(bg="white")
    master.title(title)
    master.geometry("800x600")

    e=Entry(master,width=45, borderwidth=5)
    e.pack(side=BOTTOM)#place(x=ex,y=ey)
    Label(master,text=level, font=("arial",size), bg="white").place(x=xx,y=yy)
    Label(master, text= question, bg="white", font=("arial", sizepx), justify="center").place(x=x, y=y)
    bottom = Label(master, text="",bg= "white")
    bottom.place(x=ex+50, y=ey-100)

    # --create the menus
    main_menu = Menu(master)
    master.config(menu=main_menu)

    second_menu = Menu(main_menu, tearoff=0)
    main_menu.add_cascade(label="download as docx", menu=second_menu)
    third_menu = Menu(main_menu, tearoff=0)
    main_menu.add_cascade(label="download as txt", menu=third_menu)

    question_menu_word = Menu(second_menu, tearoff=0)
    clues_menu_word = Menu(second_menu,tearoff=0)
    answer_menu_word = Menu(second_menu,tearoff=0)
    second_menu.add_cascade(label="only questions", menu=question_menu_word)
    second_menu.add_separator()
    second_menu.add_cascade(label="only clues", menu=clues_menu_word)
    second_menu.add_separator()
    second_menu.add_cascade(label="only answers", menu=answer_menu_word)
    second_menu.add_separator()

    second_menu.add_command(label="all together", command=None)

    question_menu_word.add_command(label="question 1", command = lambda: word_download(False,False, 1, [1]))
    question_menu_word.add_command(label="question 2", command = lambda: word_download(False,False, 2, [2]))
    question_menu_word.add_command(label="question 3", command = lambda: word_download(False,False, 3, [3]))
    question_menu_word.add_command(label="question 4", command = lambda: word_download(False,False, 4, [4]))
    question_menu_word.add_command(label="question 5", command = lambda: word_download(False,False, 5, [5]))
    question_menu_word.add_separator()
    question_menu_word.add_command(label="all the questions", command = lambda: word_download(False, False, 1, [1,2,3,4,5]))


    clues_menu_word.add_command(label="question 1", command = lambda: word_download(True,False, 1, [1]))
    clues_menu_word.add_command(label="question 2", command = lambda: word_download(True,False, 2, [2]))
    clues_menu_word.add_command(label="question 3", command = lambda: word_download(True,False, 3, [3]))
    clues_menu_word.add_command(label="question 4", command = lambda: word_download(True,False, 4, [4]))
    clues_menu_word.add_command(label="question 5", command = lambda: word_download(True,False, 5, [5]))
    clues_menu_word.add_separator()
    clues_menu_word.add_command(label="all the questions", command = lambda: word_download(True,False, 1, [1,2,3,4,5]))

    answer_menu_word.add_command(label="question 1", command = lambda: word_download(False,True, 1, [1]))
    answer_menu_word.add_command(label="question 2", command = lambda: word_download(False,True, 2, [2]))
    answer_menu_word.add_command(label="question 3", command = lambda: word_download(False,True, 3, [3]))
    answer_menu_word.add_command(label="question 4", command = lambda: word_download(False,True, 4, [4]))
    answer_menu_word.add_command(label="question 5", command = lambda: word_download(False,True, 5, [5]))
    answer_menu_word.add_separator()
    answer_menu_word.add_command(label="all the questions", command = lambda: word_download(False,True, 1, [1,2,3,4,5]))


    question_menu_txt = Menu(third_menu)
    clues_menu_txt = Menu(third_menu)
    answer_menu_txt = Menu(third_menu)
    third_menu.add_cascade(label="only questions", menu=question_menu_txt)
    third_menu.add_separator()
    third_menu.add_cascade(label="only clues", menu=question_menu_txt)
    third_menu.add_separator()
    third_menu.add_cascade(label="only answers", menu=question_menu_txt)
    third_menu.add_separator()

    third_menu.add_command(label="all together", command=None)


    def submit(value,result, label=None, tk=None):
        if str(value) == str(result) and level[-1] == "1": tk.destroy(),questions(setting_btn2)
        if str(value) == str(result) and level[-1] == "2": tk.destroy(),questions(setting_btn3)
        if str(value) == str(result) and level[-1] == "3": tk.destroy(),questions(setting_btn4)
        if str(value) == str(result) and level[-1] == "4": tk.destroy(),questions(setting_btn5)
        if str(value) == str(result) and level[-1] == "5": tk.destroy(),questions(setting_btn2)
        else: bottom.config(text="התשובה לא היתה נכונה")

    button = Button(master, text="submit", command=lambda: submit(e.get(),result, bottom, master))
    button.place(x=ex-50,y=ey+40)

def main(*functions, statement):
    if statement:
        for function in functions:
            function()
    else:
        debugger()

canvas.create_text(500,50,text="חדר בריחה",font=("arial",40))

question1 = """הגעת לחדר חשוך עם שתי דלתות יציאה.
 אם תצליח לפתור את החידה הבאה תוך 30 דקות
 הדלת הראשונה תפתח
 אם לא הצלחת הדלת השניה תפתח ותוביל אותך אל מסך הבית: 
הגעת לחדר חשוך בעל משאבים מוגבלים.
הטמפרטורה ההתחלתית של החדר היא אפס מעלות
כל דקה שאתה נמצא בחדר הטמפרטורה עולה ב5  מעלות
כדי לעצור את זה ולפתוח את החדר
תצטרך להדליק אש לרשותך מצית ומעבדה.

"""
setting_btn1 = (
    "question 1", "level 1", 
    question1, 20, 20, 
    60, 50, 50, 200,
    400, 400, 7)

setting_btn2 = (
    "question 2", "level 2",
    question1, 20,
    20, 100, 100, 50, 200,
    400, 400, 15)
button1 = Button(root, text="level1", command=lambda: questions(setting_btn1)
)# button 1
canvas.create_window(500,100,window=button1)

button2 = Button(root, text="level2", command=lambda: questions(setting_btn2)
)
canvas.create_window(500,200,window=button2)

button3 = Button(root, text="level3", command=lambda: questions(
    "question 3", "level 3", question1, 20,
    20, 100, 100 , 50, 200, 400 , 400, 15)
)
canvas.create_window(500,300,window=button3)

button4 = Button(root, text="level4", command=lambda: questions(
    "question 4", "level 4", question1, 20,
    20, 100, 100 , 50, 200, 400 , 400, 15)
)
canvas.create_window(500,400,window=button4)

button5 = Button(root, text="level5", command=lambda: questions(
    "question 4", "level 4", question1, 20,
    20, 100, 100 , 50, 200, 400 , 400, 15)
)
canvas.create_window(500,500,window=button5)

root.mainloop()
print(0b1010)